from itertools import count
import os
import shutil #This helps move, copy, etc.
import chardet
from docx import Document



#Lists: (Learning Note: Appending to a list within a function still saves it in the list even for other functions)
list_fallback_encodings = [
    'utf-8', 'latin-1', 'utf-16', 'utf-32', 'cp1252', 'iso-8859-1',
    'iso-8859-2', 'iso-8859-3', 'iso-8859-4', 'iso-8859-5', 'iso-8859-6',
    'iso-8859-7', 'iso-8859-8', 'iso-8859-9', 'iso-8859-10', 'iso-8859-13',
    'iso-8859-14', 'iso-8859-15', 'iso-8859-16', 'windows-1250', 'windows-1251',
    'windows-1252', 'windows-1253', 'windows-1254', 'windows-1255', 'windows-1256',
    'windows-1257', 'windows-1258', 'macintosh', 'mac-roman', 'ascii',
    'big5', 'euc-jp', 'euc-kr', 'gb2312', 'gbk', 'hz-gb-2312', 'shift-jis',
    'ks_c_5601-1987'
]

list_expected_languages = ['en', 'fr', 'de', 'es', 'ko']  # Update with the languages you expect. 
    # En = English,
    # Fr = French,
    # De = German,
    # Es = Spanish,
    # Ko = Korean,


list_encode_these_endswith_filetypes = [
    '.txt',
]

list_correct_endswith_files =[]
list_incorrect_endswith_files =[]

list_successful_encoding_files= []
list_failed_encoding_files = []



# Set the folder pathways [MANUAL]:
  # start_broken_folder = Start Folder
  # destination_resolved_folder = Destination Folder
  # could_not_resolve_folder = Any files that could not be encoded to be moved here.
start_broken_folder = '/Users/simyoonchan/Documents/EncodingProject/BrokenRound1-Test'
destination_resolved_folder = '/Users/simyoonchan/Documents/EncodingProject/ResolvedRound1'
incorrect_endswith_dumpster_folder = '/Users/simyoonchan/Documents/EncodingProject/IncorrectEndswithDumpsterFolder'
failed_encoding_files_dumpster_folder = '/Users/simyoonchan/Documents/EncodingProject/FailedEncodingFiles'



#Function Line Break:
def print_empty_row():
    print()



#Function Vertical Line: To divide major sections
def print_vertical_line():
    print("|"*80)



#Function Horizontal Line: To divide minor sections
def print_horizontal_line():
    print("-"*100)



#Function Section Separator:
def section_separator():
    print_empty_row()
    print_empty_row()
    print_vertical_line()
    


# Function:
def count_and_copy_files():
   #Initialize the counter for all files.
   files_count = 0
   

   #Step: Loop through files & Count files
   for root_folder, throwaway_variable, list_of_file_names in os.walk(start_broken_folder):
       for file in list_of_file_names:
          start_file_path = os.path.join(root_folder, file)

          #Identify file for copy.
          print(f"File ready for copy: {file} from {start_file_path}")
          print_empty_row()

          #Increment file count.
          files_count += 1

    
   #Step: Declare total number of files
   print(f"You have this number of files: {files_count}")


   #Step: Decider point
   while True: #While true loop will continuosuly prompt the user until valid response.
        user_response = input("Would you like to copy all these files and its folders? (yes copy/no):")
   
        if user_response.lower() in ['yes copy']:
            
            #Step: Copy files
                #Note: Before copying the files, you can remove the destination folder if it already exists. You can do this by using the shutil.rmtree() function to delete the destination folder and its contents. After that, you can proceed with copying the files.
            if os.path.exists(destination_resolved_folder):
                 shutil.rmtree(destination_resolved_folder)
            shutil.copytree(start_broken_folder, destination_resolved_folder)
            
            #Section separator.
            section_separator()

            #Next Function.
            move_onto_next_function_separate_endswith_files()
            break

        elif user_response.lower() in ['no']:
            #Section separator.
            section_separator()
            break

        else:
            print("Invalid input. Print 'yes copy' or 'no'.")



#Function:
def move_onto_next_function_separate_endswith_files():
    while True: #While true loop will continuosuly prompt the user until valid response.
        user_response = input("Do you want to move onto the next function? (yes move on/no):")

        if user_response.lower() in ['yes move on']:          
            #Section separator.
            section_separator()

            #Next Function.
            separate_endswith_files()
            break
        
        elif user_response.lower() in ['no']:
            #Section separator.
            section_separator()
            break

        else:
            print("Invalid input. Please enter 'yes move on' or 'no'.")



#Function:
def separate_endswith_files():
    for root_folder, throwaway_variable, list_of_file_names in os.walk(destination_resolved_folder):
        for file in list_of_file_names:
            start_file_path = os.path.join(root_folder, file)
            could_not_resolve_file_path = os.path.join(incorrect_endswith_dumpster_folder, file)

            for endswith in list_encode_these_endswith_filetypes:
                if file.endswith(endswith):
                    list_correct_endswith_files.append(file)
                    #print(f"Resolvable Identified:{file}")
                    #print_empty_row()
            
                else:
                    shutil.move(start_file_path, could_not_resolve_file_path)
                    list_incorrect_endswith_files.append(file)
                    #print(f"Unresolvable and moved: {file}")
                    #print_empty_row()


    #Print Lists.
    print(f"Correct Endswith Files Identified: {list_correct_endswith_files}")
    print_empty_row
    print(f"Incorrect Endswith Files Identified and moved: {list_incorrect_endswith_files}")
    

    #Section separator.
    section_separator()


    #Next Function.
    move_onto_next_function_encode_files()



#Function:
def move_onto_next_function_encode_files():
    while True: #While true loop will continuosuly prompt the user until valid response.
        user_response = input("Do you want to move onto the next function? (yes move on/no):")

        if user_response.lower() in ['yes move on']:          
            #Section separator.
            section_separator()

            #Next Function.
            encode_files_as_word_doc()
            break
        
        elif user_response.lower() in ['no']:
            #Section separator.
            section_separator()
            break

        else:
            print("Invalid input. Please enter 'yes move on' or 'no'.")



#Function:
def encode_files_as_word_doc():
    for root_folder, throwaway_variable, list_of_file_names in os.walk(destination_resolved_folder):
        for file in list_of_file_names:
            #Broken file.
            broken_file_start_file_path = os.path.join(root_folder, file)

            #Temporarily get rid of extension of broken file.
            file_name_without_extension = os.path.splitext(os.path.basename(broken_file_start_file_path))[0]
            
            #Replace broken file extension with .docx
            destination_docname_file_path = f'{file_name_without_extension}.docx'


            try:
                with open(broken_file_start_file_path, 'rb') as f:
                    raw_content = f.read()
                result = chardet.detect(raw_content)
                encoding = result['encoding'] if result['encoding'] else 'utf-8'

                content = raw_content.decode(encoding)
            except Exception as e:
                print(f"Error processing {broken_file_start_file_path}: {e}")
                continue

            if content:
                doc = Document()
                doc.add_paragraph(content)

                try:
                    doc.save(destination_docname_file_path)
                except Exception as e:
                    print(f"Error saving as Word document: {e}")
                    continue

                try:
                    os.remove(broken_file_start_file_path)
                except Exception as e:
                    print(f"Error deleting original file: {e}")

                print(f'Successfully decoded and saved as {destination_docname_file_path}')
            else:
                print(f'Unable to decode {broken_file_start_file_path}')



# Function to encode files as Word documents
def encode_files_as_word_doc():
    for root_folder, _, list_of_file_names in os.walk(destination_resolved_folder):
        for file in list_of_file_names:
            #Broken file.
            broken_file_start_file_path = os.path.join(root_folder, file)

            #Temporarily get rid of extension of broken file.
            file_name_without_extension, _ = os.path.splitext(os.path.basename(broken_file_start_file_path))
            
            #Replace broken file extension with .docx
            destination_docname_file_path = f'{file_name_without_extension}.docx'

            try:
                with open(broken_file_start_file_path, 'rb') as f:
                    raw_content = f.read()
                result = chardet.detect(raw_content)
                encoding = result['encoding'] if result['encoding'] else 'utf-8'

                content = raw_content.decode(encoding)
            except Exception as e:
                print(f"Error processing {broken_file_start_file_path}: {e}")
                continue

            if content:
                doc = Document()
                doc.add_paragraph(content)

                try:
                    doc.save(destination_docname_file_path)
                except Exception as e:
                    print(f"Error saving as Word document: {e}")
                    continue

                try:
                    os.remove(broken_file_start_file_path)
                except Exception as e:
                    print(f"Error deleting original file: {e}")

                print(f'Successfully decoded and saved as {destination_docname_file_path}')
            else:
                print(f'Unable to decode {broken_file_start_file_path}')


#Call Functions:
#count_and_copy_files()
encode_files_as_word_doc()


#What to input into Command Prompt in MacBook:
    #1) cd /Users/simyoonchan/Documents/EncodingProject
    #2) python3 encode_files_edit2.py 

